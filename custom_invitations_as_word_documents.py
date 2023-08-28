#!/usr/bin/env python3
# custom_invitations_as_word_documents.py — An exercise in manipulating Word documents.
# For more information, see README.md

import logging
import docx

logging.basicConfig(
    level=logging.DEBUG,
    filename="logging.txt",
    format="%(asctime)s -  %(levelname)s -  %(message)s",
)
logging.disable(logging.CRITICAL)  # Note out to enable logging.


def make_invitations():
    """Open styled blank docx, write invitation for
    each guest on separate page and save file."""

    with open("guests.txt", "r", encoding="utf-8") as f:
        guests = [guest.strip() for guest in f.readlines()]

    doc = docx.Document("template.docx")

    for idx, guest in enumerate(guests):
        doc.add_paragraph(
            "It would be a pleasure to have the company of", style="lines135"
        )
        doc.add_paragraph(guest, style="guest")
        doc.add_paragraph("at 11010 Memory Lane on the Evening of", style="lines135")
        doc.add_paragraph("April 1st", style="date")
        doc.add_paragraph("at 7 o'clock", style="lines135")
        if idx != len(guests) - 1:
            doc.add_page_break()

    doc.save("invitations.docx")


def main():
    make_invitations()


if __name__ == "__main__":
    main()
